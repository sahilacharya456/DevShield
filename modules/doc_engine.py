"""
DevShield AI — Module 2: Documentation Engine
Generates professional README, API docs, usage examples,
and exports to Word (.docx) and PDF.
"""

import json
import re
from datetime import datetime
from pathlib import Path

import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from config.settings import APP_NAME, APP_OWNER, EXPORTS_DIR, GEMINI_API_KEY, GEMINI_MODEL

genai.configure(api_key=GEMINI_API_KEY)


def _get_model() -> genai.GenerativeModel:
    return genai.GenerativeModel(model_name=GEMINI_MODEL)


# ─── Generate Documentation ───────────────────────────────────────────────────

def generate_documentation(
    code: str,
    task: str,
    language: str,
    security_report: dict = None,
) -> dict:
    """
    Generate full professional documentation for a piece of code.

    Returns:
        dict with: readme, function_docs, usage_examples, dependencies,
                   security_notes, changelog, success, generated_at
    """
    model = _get_model()

    sec_ctx = ""
    if security_report:
        sec_ctx = (
            f"\nSecurity Analysis Context:\n"
            f"- Score: {security_report.get('overall_score', 'N/A')}/100 "
            f"(Grade: {security_report.get('grade', 'N/A')})\n"
            f"- Issues Found: {security_report.get('total_issues', 0)}\n"
            f"- Summary: {security_report.get('summary', '')[:200]}\n"
        )

    prompt = f"""You are a senior technical writer. Generate comprehensive, professional documentation.

Task: {task}
Language: {language}
{sec_ctx}

Source Code:
```{language.lower().split()[0]}
{code[:4000]}{'...' if len(code) > 4000 else ''}
```

Generate ALL of the following sections:

1. README — complete Markdown with badges, overview, installation, usage, configuration, API reference, contributing guidelines
2. Function/class documentation — for EVERY function or class found
3. Three working usage examples with different scenarios
4. Dependencies list with versions and purposes
5. Security notes based on code patterns
6. Initial changelog entry

Return ONLY valid JSON (no markdown fences):
{{
  "readme": "<full markdown README>",
  "function_docs": [
    {{
      "name": "<function or class name>",
      "signature": "<full signature>",
      "description": "<what it does>",
      "params": [{{"name": "<param>", "type": "<type>", "description": "<desc>"}}],
      "returns": "<return description>",
      "raises": ["<ExceptionType>: <when>"],
      "example": "<one-liner usage>"
    }}
  ],
  "usage_examples": [
    {{
      "title": "<example title>",
      "description": "<what this demonstrates>",
      "code": "<working code>"
    }}
  ],
  "dependencies": [{{"name": "<pkg>", "version": ">=x.y", "purpose": "<why>"}}],
  "security_notes": ["<note1>", "<note2>"],
  "changelog": "## v1.0.0\\n- Initial release"
}}"""

    try:
        response = model.generate_content(prompt)
        raw = response.text.strip()
        raw = re.sub(r"^```(?:json)?\s*\n?", "", raw, flags=re.MULTILINE)
        raw = re.sub(r"\n?```\s*$", "", raw, flags=re.MULTILINE)

        data = json.loads(raw)
        data["success"] = True
        data["generated_at"] = datetime.now().isoformat()
        return data

    except Exception as exc:
        return {
            "readme": f"# {task}\n\nDocumentation generation failed: {exc}",
            "function_docs": [],
            "usage_examples": [],
            "dependencies": [],
            "security_notes": [],
            "changelog": "## v1.0.0\n- Initial release",
            "success": False,
            "error": str(exc),
            "generated_at": datetime.now().isoformat(),
        }


# ─── Word Export ──────────────────────────────────────────────────────────────

def export_to_word(
    doc_data: dict,
    task: str,
    language: str,
    code: str,
    security_report: dict = None,
    filename: str = None,
) -> str:
    """Export full documentation + code + security report to a styled Word document."""
    if not filename:
        safe = re.sub(r"[^\w\s-]", "", task)[:30].strip().replace(" ", "_")
        filename = f"DevShield_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    filepath = EXPORTS_DIR / filename
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Cover ───────────────────────────────────────────────────────────────
    title = doc.add_heading("🛡️ DevShield AI — Technical Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"Generated by {APP_NAME} for {APP_OWNER}  •  {datetime.now().strftime('%B %d, %Y')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    # Info table
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Table Grid"
    cells_data = [("Task", task), ("Language", language),
                  ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))]
    for i, (k, v) in enumerate(cells_data):
        row = tbl.rows[i]
        row.cells[0].text = k
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = v
    doc.add_paragraph()

    # ── README ──────────────────────────────────────────────────────────────
    doc.add_heading("📋 README", level=1)
    for line in doc_data.get("readme", "").split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=2)
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            p = doc.add_paragraph(clean)
            p.style.font.size = Pt(10)

    # ── Source Code ──────────────────────────────────────────────────────────
    doc.add_heading("💻 Source Code", level=1)
    code_p = doc.add_paragraph(code)
    run = code_p.runs[0]
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    # ── API Reference ────────────────────────────────────────────────────────
    fn_docs = doc_data.get("function_docs", [])
    if fn_docs:
        doc.add_heading("📚 API Reference", level=1)
        for fn in fn_docs:
            doc.add_heading(fn.get("name", "Unknown"), level=2)
            doc.add_paragraph(f"Signature: {fn.get('signature', '')}")
            doc.add_paragraph(fn.get("description", ""))
            if fn.get("params"):
                doc.add_heading("Parameters", level=3)
                for p in fn["params"]:
                    doc.add_paragraph(
                        f"• {p.get('name')} ({p.get('type', 'Any')}): {p.get('description', '')}"
                    )
            if fn.get("returns"):
                doc.add_paragraph(f"Returns: {fn['returns']}")
            if fn.get("example"):
                ep = doc.add_paragraph(fn["example"])
                ep.runs[0].font.name = "Courier New"
                ep.runs[0].font.size = Pt(9)

    # ── Usage Examples ───────────────────────────────────────────────────────
    examples = doc_data.get("usage_examples", [])
    if examples:
        doc.add_heading("🚀 Usage Examples", level=1)
        for i, ex in enumerate(examples, 1):
            doc.add_heading(f"{i}. {ex.get('title', f'Example {i}')}", level=2)
            doc.add_paragraph(ex.get("description", ""))
            ep = doc.add_paragraph(ex.get("code", ""))
            ep.runs[0].font.name = "Courier New"
            ep.runs[0].font.size = Pt(9)

    # ── Security Report ──────────────────────────────────────────────────────
    if security_report:
        doc.add_heading("🔐 Security Analysis Report", level=1)
        doc.add_paragraph(
            f"Score: {security_report.get('overall_score', 'N/A')}/100  "
            f"(Grade: {security_report.get('grade', 'N/A')})"
        ).runs[0].bold = True
        doc.add_paragraph(security_report.get("summary", ""))

        vulns = security_report.get("vulnerabilities", [])
        if vulns:
            doc.add_heading("Vulnerabilities Found", level=2)
            for v in vulns:
                doc.add_paragraph(
                    f"[{v.get('severity')}] {v.get('name')} — Line {v.get('line')}"
                ).runs[0].bold = True
                doc.add_paragraph(f"  {v.get('description', '')}")
                doc.add_paragraph(
                    f"  OWASP: {v.get('owasp_id', 'N/A')} — {v.get('owasp_name', 'N/A')}"
                )
                doc.add_paragraph(f"  Fix: {v.get('fix_suggestion', 'See description')}")
                doc.add_paragraph()

    doc.save(str(filepath))
    return str(filepath)


# ─── PDF Export ───────────────────────────────────────────────────────────────

def export_to_pdf(
    doc_data: dict,
    task: str,
    language: str,
    code: str,
    security_report: dict = None,
    filename: str = None,
) -> str:
    """Export full documentation + security report to a styled PDF."""
    if not filename:
        safe = re.sub(r"[^\w\s-]", "", task)[:30].strip().replace(" ", "_")
        filename = f"DevShield_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

    filepath = EXPORTS_DIR / filename
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "DSTitle",
        parent=styles["Title"],
        fontSize=22,
        textColor=colors.HexColor("#00CC6A"),
        spaceAfter=10,
        fontName="Helvetica-Bold",
    )
    h1_style = ParagraphStyle(
        "DSH1",
        parent=styles["Heading1"],
        fontSize=14,
        textColor=colors.HexColor("#0088CC"),
        spaceBefore=14,
        spaceAfter=6,
        fontName="Helvetica-Bold",
    )
    code_style = ParagraphStyle(
        "DSCode",
        parent=styles["Code"],
        fontSize=7.5,
        fontName="Courier",
        leftIndent=12,
        spaceAfter=4,
    )
    normal = styles["Normal"]

    story = []

    # Title
    story.append(Paragraph("&#127697; DevShield AI — Technical Report", title_style))
    story.append(
        Paragraph(
            f"Generated for {APP_OWNER}  •  {datetime.now().strftime('%B %d, %Y')}",
            normal,
        )
    )
    story.append(Spacer(1, 0.2 * inch))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#00CC6A")))
    story.append(Spacer(1, 0.15 * inch))

    # Info table
    info = [
        ["Task", task[:80]],
        ["Language", language],
        ["Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
    ]
    if security_report:
        info.append(
            [
                "Security Score",
                f"{security_report.get('overall_score', 'N/A')}/100  "
                f"(Grade {security_report.get('grade', 'N/A')})",
            ]
        )

    t = Table(info, colWidths=[1.5 * inch, 5 * inch])
    t.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.HexColor("#F4F4F4"), colors.white]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("PADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(t)
    story.append(Spacer(1, 0.2 * inch))

    # README
    readme = doc_data.get("readme", "")
    if readme:
        story.append(Paragraph("README", h1_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CCCCCC")))
        for line in readme.split("\n")[:60]:
            clean = re.sub(r"[#*`]", "", line).strip()
            if clean:
                story.append(Paragraph(clean[:200], normal))
        story.append(Spacer(1, 0.15 * inch))

    # Source code
    story.append(Paragraph("Source Code", h1_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CCCCCC")))
    display_code = code if len(code) < 3000 else code[:3000] + "\n... [truncated]"
    story.append(Preformatted(display_code, code_style))
    story.append(Spacer(1, 0.2 * inch))

    # Security Report
    if security_report:
        story.append(Paragraph("Security Analysis Report", h1_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CCCCCC")))
        story.append(
            Paragraph(
                f"Score: {security_report.get('overall_score')}/100  |  "
                f"Grade: {security_report.get('grade')}  |  "
                f"Issues: {security_report.get('total_issues', 0)}",
                normal,
            )
        )
        story.append(Paragraph(security_report.get("summary", ""), normal))
        story.append(Spacer(1, 0.1 * inch))

        vulns = security_report.get("vulnerabilities", [])
        if vulns:
            vuln_data = [["Severity", "Vulnerability", "Line", "OWASP"]]
            for v in vulns[:25]:
                vuln_data.append(
                    [
                        v.get("severity", ""),
                        v.get("name", "")[:38],
                        str(v.get("line", "")),
                        v.get("owasp_id", ""),
                    ]
                )
            t2 = Table(vuln_data, colWidths=[1 * inch, 3.2 * inch, 0.6 * inch, 1.2 * inch])
            t2.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                        ("ROWBACKGROUNDS", (1, 0), (-1, -1), [colors.HexColor("#FFF5F5"), colors.white]),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("PADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(t2)

    pdf_doc = SimpleDocTemplate(
        str(filepath),
        pagesize=letter,
        rightMargin=72, leftMargin=72,
        topMargin=72, bottomMargin=72,
    )
    pdf_doc.build(story)
    return str(filepath)
